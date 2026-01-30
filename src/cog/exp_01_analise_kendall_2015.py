"""
================================================================================
PROJECT:         Geography of Cognition: Poverty, Wealth, and Inequalities in Brazil
SCRIPT:          src/cog/exp_01_analise_kendall_2015.py
VERSION:         5.1 (Deep Statistics: PCA + Shapiro + Scientific Header)
DATE:            2026-01-29
--------------------------------------------------------------------------------
PRINCIPAL INVESTIGATOR:  Dr. José Aparecido da Silva
LEAD DATA SCIENTIST:     Me. Cássio Dalbem Barth
SOURCE:                  Internal Processed Data (PISA, SAEB, ENEM - 2015)
================================================================================
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from mpl_toolkits.mplot3d import Axes3D
import seaborn as sns
from pandas.plotting import parallel_coordinates
import os
import sys
import time
from pathlib import Path
from datetime import datetime
from scipy.stats import chi2, spearmanr, pearsonr, shapiro

# Tenta importar SKLEARN para PCA (Estatística Avançada)
try:
    from sklearn.decomposition import PCA
    from sklearn.preprocessing import StandardScaler
    HAS_SKLEARN = True
except ImportError:
    HAS_SKLEARN = False

# Tenta importar DOCX
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# --- CONFIGURAÇÃO ---
TARGET_YEAR = "2015"
BASE_DIR = Path(__file__).resolve().parent.parent.parent 
DIRS = {
    "csv":      BASE_DIR / "data" / "processed", / 'testes',
    "xlsx":     BASE_DIR / "reports" / "varcog" / "xlsx",
    "graficos": BASE_DIR / "reports" / "varcog" / "graficos",
    "relatorios": BASE_DIR / "reports" / "varcog" / "relatorios"
}
for p in DIRS.values(): p.mkdir(parents=True, exist_ok=True)

REGION_MAP = {
    'RO': 'Norte', 'AC': 'Norte', 'AM': 'Norte', 'RR': 'Norte', 'PA': 'Norte', 'AP': 'Norte', 'TO': 'Norte',
    'MA': 'Nordeste', 'PI': 'Nordeste', 'CE': 'Nordeste', 'RN': 'Nordeste', 'PB': 'Nordeste', 'PE': 'Nordeste', 'AL': 'Nordeste', 'SE': 'Nordeste', 'BA': 'Nordeste',
    'MG': 'Sudeste', 'ES': 'Sudeste', 'RJ': 'Sudeste', 'SP': 'Sudeste',
    'PR': 'Sul', 'SC': 'Sul', 'RS': 'Sul',
    'MS': 'Centro-Oeste', 'MT': 'Centro-Oeste', 'GO': 'Centro-Oeste', 'DF': 'Centro-Oeste'
}

# --- UTILS ---
try:
    import msvcrt
    def input_timeout(prompt, timeout=10, default=''):
        print(f"{prompt} [Automático em {timeout}s]: ", end='', flush=True)
        start = time.time()
        chars = []
        while True:
            if msvcrt.kbhit():
                c = msvcrt.getwche()
                if c == '\r': print(); return "".join(chars) or default
                chars.append(c)
            if time.time() - start > timeout: print(f"\n[TIMEOUT] Usando: {default}"); return default
except:
    def input_timeout(prompt, timeout=10, default=''):
        return input(f"{prompt} [Default: {default}]: ") or default

# 1. RADAR (v5.0 Logic)
def scan_files(target_suffix):
    print(f"\n[RADAR] Buscando arquivos do ciclo {TARGET_YEAR}...")
    candidates = {"pisa": [], "saeb": [], "enem": []}
    ignore = ["venv", ".git", "__pycache__"]
    
    for root, dirs, files in os.walk(BASE_DIR):
        if any(x in root for x in ignore): continue
        for file in files:
            fname = file.lower()
            if (fname.endswith(".xlsx") or fname.endswith(".csv")) and TARGET_YEAR in fname and not fname.startswith("~$"):
                path = os.path.join(root, file)
                if "pisa" in fname: candidates["pisa"].append(path)
                elif "saeb" in fname: candidates["saeb"].append(path)
                elif "enem" in fname: candidates["enem"].append(path)
    
    final = {}
    for key, paths in candidates.items():
        if not paths:
            print(f"[ERRO] Nenhum arquivo encontrado para {key.upper()}.")
            final[key] = None; continue
        paths.sort(key=lambda x: 0 if target_suffix in os.path.basename(x).lower() else 1)
        
        # Força menu apenas se houver ambiguidade ou for ENEM (para garantir escolha de filtro)
        if len(paths) > 1:
            print(f"\n[SELEÇÃO] Arquivos para {key.upper()}:")
            for i, p in enumerate(paths):
                rec = " (*)" if target_suffix in os.path.basename(p).lower() else ""
                print(f"  [{i+1}] {os.path.basename(p)}{rec}")
            sel = input(f">> Digite o número do arquivo para {key.upper()}: ").strip()
            try: final[key] = paths[int(sel)-1]
            except: final[key] = paths[0]
        else:
            final[key] = paths[0]
            print(f"[AUTO] {key.upper()}: {os.path.basename(paths[0])}")
    return final

# 2. LOAD
def smart_load(files_dict, use_weighted=True):
    if not all(files_dict.values()): return None
    try:
        def load(p): return pd.read_excel(p) if p.endswith("xlsx") else pd.read_csv(p)
        df_p, df_s, df_e = load(files_dict["pisa"]), load(files_dict["saeb"]), load(files_dict["enem"])
        
        if 'Filtro' in df_e.columns:
            print(f"   [INFO] Filtro ENEM ativo: {df_e['Filtro'].unique()}")

        def get_col(df):
            cols = df.columns
            tags = ["Ponderada", "Weighted", "w_mean"] if use_weighted else []
            tags += ["Cognitive_Global_Mean", "Média_Geral", "Mean_General", "SAEB_General", "ENEM_General"]
            for t in tags:
                for c in cols: 
                    if t in c: return c
            for c in cols:
                if "media" in c.lower() or "mean" in c.lower(): return c
            return None

        cp, cs, ce = get_col(df_p), get_col(df_s), get_col(df_e)
        if not all([cp, cs, ce]): return None
        print(f"   Colunas: PISA={cp} | SAEB={cs} | ENEM={ce}")

        d1 = df_p[['UF', cp]].rename(columns={cp: 'PISA'})
        d2 = df_s[['UF', cs]].rename(columns={cs: 'SAEB'})
        d3 = df_e[['UF', ce]].rename(columns={ce: 'ENEM'})
        return d1.merge(d2, on='UF').merge(d3, on='UF')
    except Exception as e: print(f"[ERRO] {e}"); return None

# 3. ANÁLISE BÁSICA
def run_analysis(df):
    df['R_PISA'] = df['PISA'].rank(ascending=False)
    df['R_SAEB'] = df['SAEB'].rank(ascending=False)
    df['R_ENEM'] = df['ENEM'].rank(ascending=False)
    m = 3; n = len(df)
    df['S'] = df['R_PISA'] + df['R_SAEB'] + df['R_ENEM']
    S_var = ((df['S'] - (m*(n+1)/2))**2).sum()
    W = (12 * S_var) / (m**2 * (n**3 - n))
    
    df['Rank_Consenso'] = df['S'].rank(method='min')
    df['Desvio_Rank'] = df[['R_PISA', 'R_SAEB', 'R_ENEM']].std(axis=1)
    
    conditions = [(df['Desvio_Rank'] <= 2.0), (df['Desvio_Rank'] > 5.0)]
    df['Estabilidade'] = np.select(conditions, ['Alta', 'Baixa'], default='Média')
    df['Regiao'] = df['UF'].map(REGION_MAP)
    return df, W

# 4. ESTATÍSTICA PROFUNDA (PCA + SHAPIRO)
def run_deep_stats(df):
    stats = {}
    
    # 1. Shapiro-Wilk (Normalidade)
    # H0: A distribuição é normal. Se p < 0.05, rejeita H0 (não é normal).
    for col in ['PISA', 'SAEB', 'ENEM']:
        stat, p = shapiro(df[col])
        stats[f'shapiro_{col}'] = (stat, p)
        
    # 2. PCA (Principal Component Analysis)
    if HAS_SKLEARN:
        X = df[['PISA', 'SAEB', 'ENEM']].values
        X_std = StandardScaler().fit_transform(X) # Padroniza
        pca = PCA(n_components=1) # Queremos ver quanto 1 fator explica
        pca.fit(X_std)
        
        explained_variance = pca.explained_variance_ratio_[0]
        stats['pca_explained'] = explained_variance
        stats['pca_components'] = pca.components_[0] # Pesos de cada variável no Fator G
    else:
        stats['pca_explained'] = None

    return stats

# 5. PLOTS
def plot_results(df, W, suffix):
    print(">> Gerando gráficos...")
    sns.set(style="whitegrid")
    # Fluxo
    plt.figure(figsize=(10, 6))
    plot_df = df.copy().sort_values('S')
    q1, q2 = plot_df['S'].quantile([0.33, 0.66])
    plot_df['Grupo'] = np.select([(plot_df['S']<=q1), (plot_df['S']>q1)&(plot_df['S']<=q2)], ['Alta', 'Média'], default='Baixa')
    parallel_coordinates(plot_df[['R_PISA', 'R_SAEB', 'R_ENEM', 'Grupo']], 'Grupo', color=['#2ca02c', '#1f77b4', '#d62728'], alpha=0.8)
    plt.gca().invert_yaxis(); plt.title(f'Fluxo Hierárquico - W={W:.3f}', fontsize=14)
    plt.tight_layout(); plt.savefig(DIRS["graficos"] / f"kendall_2015_fluxo{suffix}.png", dpi=150); plt.close()
    
    # 3D
    fig = plt.figure(figsize=(10, 8)); ax = fig.add_subplot(111, projection='3d')
    xs, ys, zs = df['SAEB'], df['ENEM'], df['PISA']
    ax.scatter(xs, ys, zs, c=df['S'], cmap='RdYlGn_r', s=60, depthshade=True, edgecolors='k')
    for x, y, z, s in zip(xs, ys, zs, df['S']): ax.plot([x, x], [y, y], [zs.min()*0.95, z], 'k--', alpha=0.2, linewidth=0.5)
    for i, txt in enumerate(df['UF']): ax.text(xs[i], ys[i], zs[i], txt, fontsize=7)
    ax.set_xlabel('SAEB'); ax.set_ylabel('ENEM'); ax.set_zlabel('PISA'); ax.view_init(elev=20, azim=-45)
    plt.savefig(DIRS["graficos"] / f"kendall_2015_3d{suffix}.png", dpi=150); plt.close()

# 6. RELATÓRIO CIENTÍFICO (v5.1)
def generate_scientific_report(df, W, suffix, deep_stats):
    if not HAS_DOCX: return
    print(">> Gerando RELATÓRIO CIENTÍFICO (.docx)...")
    
    doc = Document()
    
    # --- HEADER PROFISSIONAL (Tabela Invisível) ---
    table_header = doc.add_table(rows=1, cols=2)
    table_header.autofit = True
    c1 = table_header.cell(0, 0)
    c2 = table_header.cell(0, 1)
    
    # Lado Esquerdo (Título)
    p_title = c1.paragraphs[0]
    p_title.add_run('GEOGRAFIA DA COGNIÇÃO\n').bold = True
    p_title.add_run('POBREZA, RIQUEZA E DESIGUALDADES NO BRASIL')
    
    # Lado Direito (Info)
    p_info = c2.paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_info.add_run(f'EXP-01 | {TARGET_YEAR}\n').bold = True
    p_info.add_run(f'{datetime.now().strftime("%d/%m/%Y")}')

    doc.add_paragraph("_"*50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Tabela de Metadados (Estilo Ficha Técnica)
    t_meta = doc.add_table(rows=5, cols=2); t_meta.style = 'Table Grid'
    
    def row_meta(i, k, v):
        c = t_meta.rows[i].cells
        c[0].text = k; c[0].paragraphs[0].runs[0].bold = True; c[0].width = Inches(1.8)
        c[1].text = v
    
    row_meta(0, "TIPO DE ANÁLISE:", "Validação de Constructo (Triangulação Multivariada)")
    row_meta(1, "INVESTIGADORES:", "Dr. José Aparecido da Silva / Me. Cássio Dalbem Barth")
    row_meta(2, "DATA/VERSÃO:", f"{datetime.now().strftime('%d/%m/%Y')} (v5.1 Deep Stats)")
    row_meta(3, "OBJETO:", "Triangulação PISA (OECD) x SAEB (MEC) x ENEM (INEP)")
    row_meta(4, "FONTE DE DADOS:", "Microdados Processados (Pasta: data/processed)")
    doc.add_paragraph()

    # 1. ABSTRACT ESTATÍSTICO
    doc.add_heading('1. SUMÁRIO EXECUTIVO', level=1)
    
    m, n = 3, len(df)
    chi2_val = m * (n - 1) * W
    p_val_kendall = chi2.sf(chi2_val, n - 1)
    
    # PCA Interpretation
    pca_text = "indisponível (Scikit-learn ausente)"
    if deep_stats['pca_explained']:
        expl = deep_stats['pca_explained'] * 100
        pca_text = f"demonstrou que {expl:.2f}% da variância total dos três testes é explicada por um único Fator Latente (G)"
    
    doc.add_paragraph(
        f"A análise de concordância de Kendall revelou um coeficiente W = {W:.4f} (p={p_val_kendall:.4e}), "
        f"indicando uma estrutura hierárquica significativa entre os estados brasileiros. "
        f"A Análise de Componentes Principais (PCA) {pca_text}. "
        f"Isto sugere que, apesar das diferenças metodológicas, as avaliações capturam um constructo cognitivo comum."
    )

    # 2. INFERÊNCIA ESTATÍSTICA (PROFUNDA)
    doc.add_heading('2. INFERÊNCIA ESTATÍSTICA', level=1)
    
    # 2.1 Normalidade
    doc.add_heading('2.1 Teste de Normalidade (Shapiro-Wilk)', level=2)
    doc.add_paragraph("Objetivo: Verificar a distribuição das notas para selecionar o método de correlação adequado.")
    
    t_shap = doc.add_table(rows=4, cols=3); t_shap.style = 'Table Grid'
    h = t_shap.rows[0].cells; h[0].text='Variável'; h[1].text='Estatística W'; h[2].text='P-valor (Sig.)'
    
    for i, col in enumerate(['PISA', 'SAEB', 'ENEM'], 1):
        stat, p = deep_stats[f'shapiro_{col}']
        res = "Normal (p>0.05)" if p > 0.05 else "Não-Normal (p<0.05)"
        row = t_shap.rows[i].cells
        row[0].text = col; row[1].text = f"{stat:.4f}"; row[2].text = f"{p:.4f} ({res})"
    
    # 2.2 Correlações Cruzadas
    doc.add_heading('2.2 Matriz de Convergência', level=2)
    s_ps, _ = spearmanr(df['PISA'], df['SAEB']); p_ps, _ = pearsonr(df['PISA'], df['SAEB'])
    s_pe, _ = spearmanr(df['PISA'], df['ENEM']); p_pe, _ = pearsonr(df['PISA'], df['ENEM'])
    s_se, _ = spearmanr(df['SAEB'], df['ENEM']); p_se, _ = pearsonr(df['SAEB'], df['ENEM'])
    
    t_corr = doc.add_table(rows=4, cols=4); t_corr.style = 'Table Grid'
    h = t_corr.rows[0].cells; h[0].text='Par de Testes'; h[1].text='Spearman (ρ) [Rank]'; h[2].text='Pearson (r) [Linear]'; h[3].text='Coef. Determinação (R²)'
    
    def fill_corr(idx, name, s, p):
        r = t_corr.rows[idx].cells
        r[0].text=name; r[1].text=f"{s:.3f}"; r[2].text=f"{p:.3f}"; r[3].text=f"{(p**2):.3f}"
        
    fill_corr(1, 'PISA x SAEB', s_ps, p_ps)
    fill_corr(2, 'PISA x ENEM', s_pe, p_pe)
    fill_corr(3, 'SAEB x ENEM', s_se, p_se)
    
    # 3. EVIDÊNCIAS VISUAIS
    doc.add_heading('3. EVIDÊNCIAS VISUAIS', level=1)
    if (DIRS["graficos"]/f"kendall_2015_fluxo{suffix}.png").exists():
        doc.add_picture(str(DIRS["graficos"]/f"kendall_2015_fluxo{suffix}.png"), width=Inches(6))
        doc.add_paragraph("Figura 1. Fluxo Hierárquico de Rankings Estaduais.", style='Caption')
    
    if (DIRS["graficos"]/f"kendall_2015_3d{suffix}.png").exists():
        doc.add_picture(str(DIRS["graficos"]/f"kendall_2015_3d{suffix}.png"), width=Inches(5.5))
        doc.add_paragraph("Figura 2. Espaço Cognitivo Tridimensional (X=SAEB, Y=ENEM, Z=PISA).", style='Caption')
    
    doc.add_page_break()
    
    # 4. TABELA MESTRA
    doc.add_heading('4. DIAGNÓSTICO ESTADUAL (MASTER TABLE)', level=1)
    doc.add_paragraph("Dados consolidados e hierarquizados por consistência global.")
    
    t_full = doc.add_table(rows=1, cols=6); t_full.style = 'Table Grid'
    h = t_full.rows[0].cells
    h[0].text='UF'; h[1].text='PISA'; h[2].text='SAEB'; h[3].text='ENEM'; h[4].text='Consenso'; h[5].text='Estabilidade'
    
    for _, row in df.sort_values('Rank_Consenso').iterrows():
        c = t_full.add_row().cells
        c[0].text = row['UF']
        c[1].text = f"{row['PISA']:.1f} (#{int(row['R_PISA'])})"
        c[2].text = f"{row['SAEB']:.1f} (#{int(row['R_SAEB'])})"
        c[3].text = f"{row['ENEM']:.1f} (#{int(row['R_ENEM'])})"
        c[4].text = f"#{int(row['Rank_Consenso'])}"
        c[5].text = row['Estabilidade']

    out_file = DIRS["relatorios"] / f"FICHA_EXP_01_VALIDACAO_2015_FULL{suffix}.docx"
    doc.save(out_file)
    print(f"[WORD] Ficha Científica salva em: {out_file}")

# MAIN
if __name__ == "__main__":
    os.system('cls' if os.name == 'nt' else 'clear')
    print("=== ANÁLISE KENDALL 2015 (v5.1 - Deep Stats) ===")
    
    if not HAS_DOCX: print("[ERRO] Instale python-docx"); sys.exit()
    if not HAS_SKLEARN: print("[AVISO] Scikit-learn não instalado. PCA será pulado.")

    print("Escolha o método de cálculo:")
    print("1 - Média Simples")
    print("2 - Média Ponderada (Recomendado)")
    choice = input_timeout(">> Opção", default="2")
    
    is_weighted = (choice == "2")
    target_suffix = "_ponderada" if is_weighted else "_simples"
    
    files = scan_files(target_suffix)
    df_merged = smart_load(files, use_weighted=is_weighted)
    
    if df_merged is not None:
        # Roda Análises
        df_final, W = run_analysis(df_merged)
        deep_stats = run_deep_stats(df_final)
        
        # Menu Output
        print("\n[OUTPUT] Gerar Relatório?")
        out_op = input_timeout(">> (S/N)", default="S")

        if out_op.upper() == 'S':
            f_csv = DIRS["csv"] / f"kendall_final_2015{target_suffix}.csv"
            f_xlsx = DIRS["xlsx"] / f"kendall_final_2015{target_suffix}.xlsx"
            df_final.to_csv(f_csv, index=False)
            df_final.to_excel(f_xlsx, index=False, sheet_name='METRICAS_2015')
            
            plot_results(df_final, W, target_suffix)
            generate_scientific_report(df_final, W, target_suffix, deep_stats)
            
        print(f"\n[SUCESSO] Pipeline finalizado.")